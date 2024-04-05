from docx import Document
from docx.shared import Pt

# Open the existing document
doc_1 = Document("test.docx")
doc_2 = Document()

# Function to determine the number of characters to bold based on word length
def chars_to_bold(word_length):
    if word_length <= 4:
        return 1
    elif word_length <= 8:
        return 2
    elif word_length <= 10:
        return 3
    else:
        return 4

# Iterate through each paragraph in the existing document
for paragraph in doc_1.paragraphs:
    # Create a new paragraph in the new document
    new_paragraph = doc_2.add_paragraph()

    # Iterate through each run (word) in the paragraph
    for run in paragraph.runs:
        words = run.text.split()

        for word in words:
            length = len(word)
            bold_chars = chars_to_bold(length)
            # Add first part of the word in bold
            modified_run = new_paragraph.add_run(word[:bold_chars])
            modified_run.bold = True
            # Add the rest of the word 
            if bold_chars < length:
                new_paragraph.add_run(word[bold_chars:])
            # Add space after the word
            new_paragraph.add_run(" ")
        
        new_paragraph.style.font.size = run.font.size
        new_paragraph.style.font.name = run.font.name

# Save the modified document
doc_2.save("test_modified.docx")
print("Text modification complete. Modified document saved as 'test_modified.docx'.")
